import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client=Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"


job_description="""
The Enterprise Technology Services organization partners with every part of the American Express business to power the company’s growth and innovation with trust and efficiency, and drive competitive differentiation with speed. We support the delivery and operations of technology, digital, and data capabilities, platforms, and services globally. Specifically, our team is responsible for the company’s technology engineering, architecture, and infrastructure, providing 24x7 support to ensure an uninterrupted, high-quality experience for customers and colleagues. We also provide product management for core enterprise platforms, and lead technology risk and information security, enterprise data governance and platforms, digital product and design, and enterprise AI platforms on behalf of the company.

American Express is embarking on an exciting transformation driven by an energetic new team of high performers. This group is nimble and creative with the power to shape our technology and product roadmap. If you have the talent and desire to deliver innovative digital and servicing products at a rapid pace, serving our customers seamlessly across physical, digital, mobile, and social media, join our transformation team! You will be part of a fast-paced, entrepreneurial team responsible for delivering projects platform supporting our global customer base. 

Our Engineers that join our Technologies team will be assigned to one of several exciting teams that are responsible for development and management of business critical platforms. You will be part of a fast-paced, entrepreneurial team responsible for delivering projects platform supporting our global customer base. You will be challenged with identifying innovative ideas and proof of concept to deliver against the existing and future needs of our customers. Our Software Engineers not only understand how technology works, but how that technology intersects with the people who count on it every day. Today, innovative ideas, insight and new perspectives are at the core of how we create a more powerful, personal and fulfilling experience for all our customers.

Responsibilities
Develop, test, debug, and maintain software applications and services under the guidance of senior engineers. 
Participate in user story analysis, sprint planning, and design discussions to understand business and technical requirements. 
Write clean, maintainable, and well-documented code following established engineering standards and best practices. 
Assist in developing RESTful APIs, microservices, and application components using modern development frameworks and tools. 
Participate in code reviews, unit testing, integration testing, and defect resolution to ensure high-quality software delivery. 
Collaborate with cross-functional teams, including Product, QA, and Engineering, to deliver features and enhancements in an Agile environment. 
Support the automation of testing, deployment, and operational processes to improve development efficiency. 
Troubleshoot and resolve application issues under the guidance of senior team members. 
Learn and adopt new technologies, development frameworks, and engineering best practices to continuously improve technical skills. 
Contribute to technical documentation, knowledge sharing, and team collaboration to support ongoing development activities.
Qualifications
Functional Skills

Strong analytical and problem-solving skills. 
Ability to understand and interpret business and technical requirements. 
Basic understanding of Agile methodologies and software development lifecycle (SDLC). 
Knowledge of object-oriented programming (OOP) concepts and software design principles. 
Hands-on experience in software development, unit testing, and debugging through academic projects, internships, or professional experience. 
Familiarity with application configuration, deployment, and version control tools (e.g., Git). 
Basic understanding of REST APIs, microservices, and distributed application concepts. 
Ability to learn new technologies quickly and adapt to changing business and technical requirements. 
Strong communication, collaboration, and teamwork skills. 
Exposure to cloud platforms (GCP, AWS, or Azure). 
Familiarity with CI/CD tools and DevOps practices. 
Basic knowledge of relational and/or NoSQL databases. 
Experience working on academic, internship, or personal software development projects.
Bachelor’s Degree in Computer Science, Computer Science Engineering, or related field required; advanced degree is preferred.
 

Preferred Qualifications

Bachelor’s Degree in Computer Science, Computer Science Engineering, or related field required; advanced degree is preferred.
6 months + years of software development experience 
Basic knowledge of cloud platforms such as Google Cloud Platform (GCP), AWS, or Azure, with exposure to services such as Google Cloud Storage and BigQuery. 
Familiarity with Big Data technologies such as Spark, Scala, Hive, or SQL 
Basic understanding of ETL (Extract, Transform, Load) concepts, data modeling, and data processing techniques. 
Exposure to relational and/or NoSQL databases and SQL programming. 
Familiarity with Agile software development methodologies and the Software Development Life Cycle (SDLC). 
Basic understanding of Continuous Integration and Continuous Delivery (CI/CD) concepts and tools such as Git, Jenkins, Maven, or similar technologies. 
Working knowledge of version control systems (Git) and build automation tools. 
Basic understanding of web technologies (HTTP, REST APIs) and familiarity with Linux/Unix environments. 
Strong understanding of programming fundamentals, data structures, algorithms, and object-oriented programming concepts. 
Strong analytical, problem-solving, and debugging skills. 
Excellent verbal and written communication skills, with the ability to work effectively in a collaborative team environment. 
Demonstrated curiosity, learning agility, and a proactive mindset toward continuous learning and process improvement.
 

Secondary Skills

Basic knowledge of Power BI for developing reports, dashboards, and data visualizations. 
Familiarity with creating simple data models, metrics, and KPIs to support reporting and business analytics. 
Working knowledge of DAX, Power Query, and SQL, with the ability to develop interactive reports and perform basic data analysis. 
Understanding of data modeling concepts and exposure to tabular models or semantic models through academic projects, internships, or professional experience.  
Basic troubleshooting and debugging skills for BI reports, datasets, and dashboards. 
Familiarity with BI concepts, metadata management, and data governance principles. 
Exposure to data visualization best practices and the ability to present insights in a clear and meaningful manner. 
Strong analytical skills with an interest in leveraging data to support business decision-making. 
Eagerness to learn modern BI technologies and continuously improve technical and analytical skills.
 

At American Express, our culture is built on a 175-year history of innovation, shared values and Leadership Behaviors, and an unwavering commitment to back our customers, communities, and colleagues. From delivering differentiated products to providing world-class customer service, we operate with a strong risk mindset, ensuring we continue to uphold our brand promise of trust, security, and service.

As part of Team Amex, you’ll experience our powerful backing with comprehensive support for your holistic well-being and many opportunities to learn new skills, develop as a leader, and grow your career. Here, your voice and ideas matter, your work makes an impact, and together, you will help us define the future of American Express.

We back you with benefits that support your holistic well-being so you can be and deliver your best. This means caring for you and your loved ones' physical, financial, and mental health, as well as providing the flexibility you need to thrive personally and professionally:

Competitive base salaries
Bonus incentives
Support for financial-well-being and retirement
Comprehensive medical, dental, vision, life insurance, and disability benefits (depending on location)
Flexible working model with hybrid, onsite or virtual arrangements depending on role and business need
Generous paid parental leave policies (depending on your location)
Free access to global on-site wellness centers staffed with nurses and doctors (depending on location)
Free and confidential counseling support through our Healthy Minds program
Career development and training opportunities
American Express is an equal opportunity employer and makes employment decisions without regard to race, color, religion, sex, sexual orientation, gender identity, national origin, veteran status, disability status, age, or any other status protected by law.

Offer of employment with American Express is conditioned upon the successful completion of a background verification check, subject to applicable laws and regulations.


"""
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system={
    "role" : "system",
    "content" : system_prompt
}
message_user={
    "role" : "user",
    "content" : user_prompt
}
response_format={
    "type" : "json_object"
}


messages=[message_system, message_user]

response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)


answer=response.choices[0].message.content

raw_json=answer
# print(raw_json)



import json
job_data=json.loads(raw_json)

job = JobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)



#parse real
class MatchResult(BaseModel):
    score: float
    details: dict
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


from PyPDF2 import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None



# lets do it now
resume_folder = Path("resumes")
all_results=[]

for file_path in resume_folder.iterdir():
    #C:\Users\Pratyush\padho_with_pratyush\week1\day5\resumes\abhay resume new - Abhay Singh.pdf
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    time.sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    #score and details
    #acount chtgpt
    # request bhejna shhur krega millions
    #chattgot server jam ho jayega
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2 = all_results[-2:]


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])